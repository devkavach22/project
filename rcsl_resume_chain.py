#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
resume_extractor.py
-------------------
Reads a resume (PDF / DOCX / DOC), extracts text, sends it to an
Agno + Groq agent, and returns a validated ResumeData Pydantic model.

Usage:
    python resume_extractor.py resume.pdf   --api-key GROQ_KEY
    python resume_extractor.py resume.docx  --api-key GROQ_KEY --model llama-3.3-70b-versatile
"""

from __future__ import annotations

import argparse
import json
import logging
import os
import re
import string
import subprocess
import sys
import tempfile
from typing import Any, List, Optional
from schemas.cv_data_schema import CVSchema as ResumeData

import pytesseract
from docx import Document as DocxDocument
from pdf2image import convert_from_path
from pdfminer.high_level import extract_text
from pydantic import BaseModel, Field, field_validator

from agno.agent import Agent
from agno.models.groq import Groq

logging.basicConfig(level=logging.INFO, format="%(levelname)s  %(message)s")
_logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = (".pdf", ".docx", ".doc")


# ============================================================
# Pydantic models
# ============================================================

def _none_str(v: Any) -> Any:
    """
    Applied as a before-validator on every Optional[str] field.
    - Converts the string literal "None"  -> Python None
    - Coerces int / float to str  (LLM sometimes returns 2 instead of "2")
    - Leaves None, real strings, and other types untouched
    """
    if v is None:
        return None
    if isinstance(v, str):
        return None if v.strip() == "None" else v
    if isinstance(v, (int, float)):
        return str(v)   # e.g. total_companies_worked: 2  ->  "2"
    return v


class ContactInfo(BaseModel):
    email:             Optional[str] = None
    phone:             Optional[str] = None
    location:          Optional[str] = None
    linkedin:          Optional[str] = None
    github:            Optional[str] = None
    portfolio:         Optional[str] = None
    address:           Optional[str] = None
    alternative_phone: Optional[str] = None

    @field_validator("*", mode="before")
    @classmethod
    def _ns(cls, v): return _none_str(v)


class Experience(BaseModel):
    company:               Optional[str]       = None
    job_title:             Optional[str]       = None
    location:              Optional[str]       = None
    start_date:            Optional[str]       = None
    end_date:              Optional[str]       = None
    duration:              Optional[str]       = None
    is_current_employment: Optional[bool]      = None
    employment_type:       Optional[str]       = None
    salary:                Optional[str]       = None
    responsibilities:      Optional[List[str]] = Field(default_factory=list)
    technologies:          Optional[List[str]] = Field(default_factory=list)

    @field_validator(
        "company", "job_title", "location", "start_date",
        "end_date", "duration", "employment_type", "salary",
        mode="before",
    )
    @classmethod
    def _ns(cls, v): return _none_str(v)


class Education(BaseModel):
    education_level: Optional[str] = None
    institution:     Optional[str] = None
    degree:          Optional[str] = None
    field_of_study:  Optional[str] = None
    course_type:     Optional[str] = None
    start_date:      Optional[str] = None
    end_date:        Optional[str] = None
    passing_year:    Optional[str] = None
    gpa:             Optional[str] = None

    @field_validator("*", mode="before")
    @classmethod
    def _ns(cls, v): return _none_str(v)


class Skills(BaseModel):
    hard_skills: Optional[List[str]] = Field(default_factory=list)
    soft_skills: Optional[List[str]] = Field(default_factory=list)


class Certification(BaseModel):
    name:            Optional[str] = None
    issuer:          Optional[str] = None
    completion_id:   Optional[str] = None
    url:             Optional[str] = None
    start_month:     Optional[str] = None
    start_year:      Optional[str] = None
    end_month:       Optional[str] = None
    end_year:        Optional[str] = None
    does_not_expire: Optional[str] = None
    date_obtained:   Optional[str] = None

    @field_validator("*", mode="before")
    @classmethod
    def _ns(cls, v): return _none_str(v)


class Language(BaseModel):
    name:        Optional[str] = None
    proficiency: Optional[str] = None

    @field_validator("*", mode="before")
    @classmethod
    def _ns(cls, v): return _none_str(v)


class Award(BaseModel):
    title:  Optional[str] = None
    issuer: Optional[str] = None
    date:   Optional[str] = None

    @field_validator("*", mode="before")
    @classmethod
    def _ns(cls, v): return _none_str(v)


class SocialMedia(BaseModel):
    name:        Optional[str] = None
    url:         Optional[str] = None
    description: Optional[str] = None

    @field_validator("*", mode="before")
    @classmethod
    def _ns(cls, v): return _none_str(v)


class WorkSample(BaseModel):
    title:       Optional[str] = None
    url:         Optional[str] = None
    description: Optional[str] = None

    @field_validator("*", mode="before")
    @classmethod
    def _ns(cls, v): return _none_str(v)


class WhitePaper(BaseModel):
    title:       Optional[str] = None
    url:         Optional[str] = None
    description: Optional[str] = None

    @field_validator("*", mode="before")
    @classmethod
    def _ns(cls, v): return _none_str(v)


class Presentation(BaseModel):
    title:       Optional[str] = None
    url:         Optional[str] = None
    description: Optional[str] = None

    @field_validator("*", mode="before")
    @classmethod
    def _ns(cls, v): return _none_str(v)


class Patent(BaseModel):
    title:       Optional[str] = None
    number:      Optional[str] = None
    date:        Optional[str] = None
    description: Optional[str] = None

    @field_validator("*", mode="before")
    @classmethod
    def _ns(cls, v): return _none_str(v)


class Project(BaseModel):
    title:        Optional[str]       = None
    description:  Optional[str]       = None
    url:          Optional[str]       = None
    technologies: Optional[List[str]] = Field(default_factory=list)
    start_date:   Optional[str]       = None
    end_date:     Optional[str]       = None

    @field_validator("title", "description", "url", "start_date", "end_date", mode="before")
    @classmethod
    def _ns(cls, v): return _none_str(v)


# class ResumeData(BaseModel):
#     full_name:                      Optional[str] = None
#     gender:                         Optional[str] = None
#     dob:                            Optional[str] = None
#     religion:                       Optional[str] = None
#     marital_status:                 Optional[str] = None
#     current_company:                Optional[str] = None
#     current_designation:            Optional[str] = None
#     total_experience:               Optional[str] = None
#     relevant_experience:            Optional[str] = None
#     total_companies_worked:         Optional[str] = None
#     current_country:                Optional[str] = None
#     current_state:                  Optional[str] = None
#     current_city:                   Optional[str] = None
#     preferred_country:              Optional[str] = None
#     preferred_state:                Optional[str] = None
#     preferred_city:                 Optional[str] = None
#     hometown:                       Optional[str] = None
#     graduation_degree:              Optional[str] = None
#     graduation_specialization:      Optional[str] = None
#     graduation_year:                Optional[str] = None
#     post_graduation_degree:         Optional[str] = None
#     post_graduation_specialization: Optional[str] = None
#     post_graduation_year:           Optional[str] = None
#     department:                     Optional[str] = None
#     role:                           Optional[str] = None
#     industry:                       Optional[str] = None
#     reason_for_change:              Optional[str] = None
#     is_permanent:                   Optional[str] = None
#     is_contractual:                 Optional[str] = None
#     is_full_time:                   Optional[str] = None
#     is_part_time:                   Optional[str] = None
#     preferred_shift:                Optional[str] = None
#     preferred_work_locations:       Optional[List[str]] = Field(default_factory=list)
#     expected_ctc:                   Optional[str] = None
#     notice_period:                  Optional[str] = None
#     status:                         Optional[str] = None
#     remarks:                        Optional[str] = None
#     summary:                        Optional[str] = None
#     contact_info:    Optional[ContactInfo]         = None
#     experience:      Optional[List[Experience]]    = Field(default_factory=list)
#     education:       Optional[List[Education]]     = Field(default_factory=list)
#     skills:          Optional[Skills]              = None
#     projects:        Optional[List[Project]]       = Field(default_factory=list)
#     certifications:  Optional[List[Certification]] = Field(default_factory=list)
#     languages:       Optional[List[Language]]      = Field(default_factory=list)
#     awards:          Optional[List[Award]]         = Field(default_factory=list)
#     social_media:    Optional[List[SocialMedia]]   = Field(default_factory=list)
#     work_samples:    Optional[List[WorkSample]]    = Field(default_factory=list)
#     white_papers:    Optional[List[WhitePaper]]    = Field(default_factory=list)
#     presentations:   Optional[List[Presentation]]  = Field(default_factory=list)
#     patents:         Optional[List[Patent]]        = Field(default_factory=list)

#     @field_validator(
#         "full_name", "gender", "dob", "religion", "marital_status",
#         "current_company", "current_designation",
#         "total_experience", "relevant_experience", "total_companies_worked",
#         "current_country", "current_state", "current_city",
#         "preferred_country", "preferred_state", "preferred_city", "hometown",
#         "graduation_degree", "graduation_specialization", "graduation_year",
#         "post_graduation_degree", "post_graduation_specialization", "post_graduation_year",
#         "department", "role", "industry", "reason_for_change",
#         "is_permanent", "is_contractual", "is_full_time", "is_part_time",
#         "preferred_shift", "expected_ctc", "notice_period",
#         "status", "remarks", "summary",
#         mode="before",
#     )
#     @classmethod
#     def _ns(cls, v): return _none_str(v)


# ============================================================
# Text extraction helpers
# ============================================================



def extract_pdf_text(file_path: str, use_ocr: bool = False) -> str:
    txt = extract_text(file_path) or ""
    if not txt or all(c in string.whitespace + "\x0c" for c in txt):
        if use_ocr:
            _logger.info("Falling back to OCR for '%s'", file_path)
            txt = _ocr_pdf(file_path)
    return txt.strip()


def _ocr_pdf(pdf_path: str) -> str:
    try:
        images = convert_from_path(pdf_path, dpi=200)
    except Exception:
        _logger.exception("Poppler conversion failed")
        return ""
    pages = []
    for i, img in enumerate(images[:5]):
        try:
            pages.append(f"--- Page {i+1} ---\n{pytesseract.image_to_string(img)}")
        except Exception:
            _logger.exception("Tesseract failed on page %s", i + 1)
    return "\n".join(pages)


def _extract_docx_text(docx_path: str) -> str:
    try:
        doc   = DocxDocument(docx_path)
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_txt = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_txt:
                    lines.append(row_txt)
        return "\n".join(lines).strip()
    except Exception:
        _logger.exception("python-docx failed")
        return ""


def _convert_doc_to_docx(doc_path: str) -> Optional[str]:
    out_dir = tempfile.mkdtemp()
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx",
             "--outdir", out_dir, doc_path],
            check=True, timeout=60,
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
        )
        converted = [
            os.path.join(out_dir, f)
            for f in os.listdir(out_dir) if f.endswith(".docx")
        ]
        return converted[0] if converted else None
    except Exception:
        _logger.exception("LibreOffice conversion failed")
        return None


def extract_text_from_file(file_path: str, use_ocr: bool = False) -> str:
    name = file_path.lower()
    if not any(name.endswith(ext) for ext in ALLOWED_EXTENSIONS):
        raise ValueError(f"Unsupported file type: {file_path}")
    if name.endswith(".pdf"):
        return extract_pdf_text(file_path, use_ocr=use_ocr)
    if name.endswith(".docx"):
        return _extract_docx_text(file_path)
    # .doc
    docx_path = _convert_doc_to_docx(file_path)
    if not docx_path:
        raise RuntimeError("LibreOffice conversion failed for .doc file")
    try:
        return _extract_docx_text(docx_path)
    finally:
        if os.path.exists(docx_path):
            os.unlink(docx_path)


# ============================================================
# Agno agent  –– fixed: no system_prompt kwarg
# ============================================================

INSTRUCTIONS = """
You are an expert resume parser.

Your task is to extract structured information from resume text with 100% completeness and accuracy.

STRICT RULES:
- Return ONLY valid JSON (no markdown, no explanation, no extra text)
- Follow the exact schema and field names provided
- Use null for missing scalar values
- Use [] for missing list fields
- NEVER skip any section, sentence, or information from the resume
- Extract ALL possible data even if it appears multiple times or in different formats
- Preserve ALL meaningful content from the resume
- Do not summarize, shorten, or ignore any information
- If a section exists (like projects, experience, skills), it MUST be extracted

CRITICAL:
- Do NOT skip any word, sentence, or character from the resume
- Ensure maximum data coverage from the CV
- Even small or incomplete sections must be extracted

Output must be a single valid JSON object only.
""".strip()

EXTRACTION_PROMPT = """\
Extract every possible field from the resume below and return a single JSON object
with EXACTLY these keys (use null when not found):

full_name, gender, dob, religion, marital_status,
current_company, current_designation,
total_experience, relevant_experience, total_companies_worked,
current_country, current_state, current_city,
preferred_country, preferred_state, preferred_city, hometown,
graduation_degree, graduation_specialization, graduation_year,
post_graduation_degree, post_graduation_specialization, post_graduation_year,
department, role, industry, reason_for_change,
is_permanent, is_contractual, is_full_time, is_part_time,
preferred_shift,
preferred_work_locations: [],
expected_ctc, notice_period, status, remarks, summary,

contact_info: {{
  email, phone, location, linkedin, github, portfolio, address, alternative_phone
}},

experience: [{{
  company, job_title, location, start_date, end_date, duration,
  is_current_employment (boolean), employment_type, salary,
  responsibilities (array of strings), technologies (array of strings)
}}],

education: [{{
  education_level, institution, field_of_study,
  course_type, passing_year, gpa
}}],

skills: {{
  hard_skills: [],
  soft_skills: []
}},

projects: [{{
  title, description, url, technologies, start_date, end_date
}}],

certifications: [{{
  name, issuer, completion_id, url,
  start_month, start_year, end_month, end_year,
  does_not_expire (boolean), date_obtained
}}],

languages: [{{
  language
}}],

awards: [{{
  title, issuer, date
}}],

social_media: [{{
  name, url, description
}}],

work_samples: [{{
  title, url, start_date, duration_month, is_currently_working, total_duration
}}],

white_papers: [{{
  title, url, start_date, duration_month, description
}}],

presentations: [{{
  title, url, description
}}],

patents: [{{
  title, url, is_issued, is_pending, application_number,
  issue_year, issue_month, description
}}]

IMPORTANT EXTRACTION RULES:

- DO NOT skip any information from the resume
- DO NOT skip any section such as Projects, Experience, Skills, Certifications
- Extract ALL projects even if they are small or one-line
- Extract ALL experience entries and responsibilities
- Extract ALL skills (technical + soft)
- Extract ALL education entries
- Extract ALL certifications, languages, awards, and other sections
- Always return string values for fields defined as string (never return arrays for string fields).
- If multiple values exist (e.g., multiple phone numbers), combine them into a single comma-separated string.
- Do NOT return lists for fields like phone, industry, role, etc.

CRITICAL RULE:
- Do NOT miss any word, sentence, or character from the resume text
- Every piece of information must be captured in the structured output
- If information cannot fit a field, include it in "remarks" or "summary"

RESUME TEXT:
{resume_text}
"""


def build_agent(api_key: str, model_id: str) -> Agent:
    """
    Build Agno Agent.
      - instructions  : system-level persona (correct kwarg, not system_prompt)
      - use_json_mode : forces the model to output valid JSON
      - max_tokens    : 8192 avoids truncated JSON for large resumes
    """
    return Agent(
        model=Groq(id=model_id, api_key=api_key, max_tokens=8192),
        instructions=INSTRUCTIONS,
        use_json_mode=True,
    )


def _repair_json(text: str) -> str:
    """
    Best-effort JSON repair for common LLM output problems:
      1. Trailing commas before ] or }
      2. Truncated JSON — close any unclosed braces/brackets
      3. Single-quoted strings  → double-quoted
      4. Bare None/True/False   → null/true/false
    """
    # 1. Replace Python literals
    text = re.sub(r'\bNone\b',  'null',  text)
    text = re.sub(r'\bTrue\b',  'true',  text)
    text = re.sub(r'\bFalse\b', 'false', text)

    # 2. Remove trailing commas before closing bracket/brace
    text = re.sub(r',\s*([\]}])', r'\1', text)

    # 3. Try to close truncated JSON by counting open braces/brackets
    open_braces   = text.count('{') - text.count('}')
    open_brackets = text.count('[') - text.count(']')
    # Close in reverse order of what's likely open last
    text = text.rstrip()
    if text.endswith(','):
        text = text[:-1]
    text += ']' * max(open_brackets, 0)
    text += '}' * max(open_braces, 0)

    return text


def _extract_content_str(response) -> str:
    """Pull a plain string out of whatever Agno returns."""
    content = response.content if hasattr(response, "content") else response

    if isinstance(content, str):
        return content

    if isinstance(content, dict):
        return json.dumps(content)

    if hasattr(content, "model_dump"):
        return json.dumps(content.model_dump(mode="json"))

    if isinstance(content, list):
        # Agno can return a list of RunResponse / message blocks
        parts = []
        for block in content:
            if isinstance(block, str):
                parts.append(block)
            elif hasattr(block, "content"):
                inner = block.content
                if isinstance(inner, str):
                    parts.append(inner)
                elif isinstance(inner, list):
                    for b in inner:
                        if hasattr(b, "text"):
                            parts.append(b.text)
            elif hasattr(block, "text"):
                parts.append(block.text)
        return "\n".join(parts)

    return str(content)


def sanitize_llm_output(data: dict) -> dict:
    """
    Fix common LLM schema issues before Pydantic validation
    """

    # --- Fix top-level string fields ---
    if "graduation_year" in data and isinstance(data["graduation_year"], int):
        data["graduation_year"] = str(data["graduation_year"])

    if "post_graduation_year" in data and isinstance(data.get("post_graduation_year"), int):
        data["post_graduation_year"] = str(data["post_graduation_year"])

    # --- Fix education ---
    for edu in data.get("education", []):
        if isinstance(edu.get("passing_year"), int):
            edu["passing_year"] = str(edu["passing_year"])

    # --- Fix experience ---
    for exp in data.get("experience", []):
        if exp.get("responsibilities") is None:
            exp["responsibilities"] = []

        if exp.get("technologies") is None:
            exp["technologies"] = []

    return data


def _parse_llm_response(response) -> dict:
    """
    Robustly parse the LLM response into a dict.
    Tries strict JSON first, then repair, then regex-extracted fragment.
    Logs the raw response for debugging when all attempts fail.
    """
    raw_str = _extract_content_str(response)
    _logger.debug("Raw LLM content:\n%s", raw_str[:1000])

    # Strip markdown fences
    cleaned = re.sub(r"```(?:json)?", "", raw_str).strip().strip("`").strip()

    if not cleaned:
        raise ValueError("LLM returned an empty response.")

    # Attempt 1 – strict parse
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError as e:
        _logger.warning("Strict JSON parse failed (%s), attempting repair…", e)

    # Attempt 2 – repair then parse
    repaired = _repair_json(cleaned)
    try:
        return json.loads(repaired)
    except json.JSONDecodeError as e:
        _logger.warning("Repaired JSON parse failed (%s), trying regex extraction…", e)

    # Attempt 3 – extract the largest {...} block and repair that
    matches = list(re.finditer(r'\{', cleaned))
    for m in reversed(matches):
        fragment = cleaned[m.start():]
        try:
            return json.loads(_repair_json(fragment))
        except json.JSONDecodeError:
            continue

    # Log the full raw output so the user can see what came back
    _logger.error("Full LLM raw output (first 2000 chars):\n%s", raw_str[:2000])
    raise ValueError(
        "Could not parse LLM response as JSON after all repair attempts. "
        "See logs for the raw output. Try switching to a larger model with "
        "--model llama-3.3-70b-versatile"
    )


# ============================================================
# Main pipeline
# ============================================================

def extract_resume(
    file_path: str,
    api_key: str,
    model_id: str = "llama-3.3-70b-versatile",
    use_ocr: bool = False,
) -> ResumeData:
    # 1 – extract text
    _logger.info("Extracting text from '%s'", file_path)
    text = extract_text_from_file(file_path, use_ocr=use_ocr)
    if not text:
        raise ValueError("No text could be extracted from the file.")
    _logger.info("Extracted %d characters", len(text))

    # 2 – send to LLM
    _logger.info("Sending to Agno agent (model: %s)", model_id)
    agent   = build_agent(api_key, model_id)
    prompt  = EXTRACTION_PROMPT.format(resume_text=text)
    response = agent.run(prompt)

    # 3 – parse response
    raw = _parse_llm_response(response)

    # ✅ FIX HERE
    raw = sanitize_llm_output(raw)

    _logger.info("LLM returned %d top-level keys", len(raw))

    # 4 – validate with Pydantic
    resume = ResumeData.model_validate(raw)
    _logger.info("Successfully parsed resume for: %s", resume.full_name)
    return resume


# ============================================================
# CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="Extract structured data from a resume (PDF/DOCX/DOC) using Agno + Groq."
    )
    parser.add_argument("file",      help="Path to resume file")
    parser.add_argument("--api-key", default=os.getenv("GROQ_API_KEY"), help="Groq API key")
    parser.add_argument("--model",   default="llama-3.3-70b-versatile", help="Groq model ID")
    parser.add_argument("--ocr",     action="store_true", help="Enable OCR for scanned PDFs")
    parser.add_argument("--output",  help="Save JSON output to this file")
    args = parser.parse_args()

    if not args.api_key:
        print(
            "[ERROR] Groq API key required. Use --api-key or set GROQ_API_KEY env var.",
            file=sys.stderr,
        )
        sys.exit(1)

    try:
        resume = extract_resume(
            file_path=args.file,
            api_key=args.api_key,
            model_id=args.model,
            use_ocr=args.ocr,
        )
        output_json = resume.model_dump_json(indent=2)
        with open("resume.json", "w") as f:
            f.write(output_json)
        print(output_json)

        if args.output:
            with open(args.output, "w", encoding="utf-8") as f:
                f.write(output_json)
            _logger.info("Saved to %s", args.output)

    except Exception as exc:
        _logger.exception("Extraction failed")
        print(f"[ERROR] {exc}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()